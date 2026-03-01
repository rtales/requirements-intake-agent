from docx import Document
from docx.shared import Pt

def _add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)

def _add_kv(doc: Document, key: str, value: str):
    p = doc.add_paragraph()
    run = p.add_run(f"{key}: ")
    run.bold = True
    p.add_run(value or "")

def render_onepager_docx(pkg: dict, out_path):
    doc = Document()
    title = doc.add_paragraph(pkg["meta"]["title"])
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)

    doc.add_paragraph("Agent: Requirements Intake Agent")

    m = pkg["meta"]
    _add_kv(doc, "ID", m["id"])
    _add_kv(doc, "Version", m.get("version","v1"))
    _add_kv(doc, "Status", m.get("status","draft"))
    _add_kv(doc, "Sensitivity", m["classification"]["sensitivity"])
    _add_kv(doc, "Domain", m["classification"]["domain"])

    _add_heading(doc, "Problem", 2)
    doc.add_paragraph(pkg["problem"]["context"])

    _add_heading(doc, "Goal", 2)
    doc.add_paragraph(pkg["problem"]["goal"])

    _add_heading(doc, "Scope", 2)
    doc.add_paragraph("In scope:")
    for x in pkg["scope"]["in_scope"]:
        doc.add_paragraph(x, style="List Bullet")
    doc.add_paragraph("Out of scope:")
    for x in pkg["scope"]["out_of_scope"]:
        doc.add_paragraph(x, style="List Bullet")

    _add_heading(doc, "Top requirements", 2)
    for r in pkg["requirements"][:8]:
        doc.add_paragraph(f"{r['id']} ({r['priority']}): {r['statement']}", style="List Number")

    _add_heading(doc, "Top acceptance criteria", 2)
    for ac in pkg["acceptance_criteria"][:8]:
        doc.add_paragraph(f"{ac['id']}: {ac['criterion']}", style="List Bullet")

    _add_heading(doc, "Top open questions", 2)
    for q in pkg["open_questions"][:8]:
        doc.add_paragraph(f"{q['id']}: {q['question']} (target: {q['target']})", style="List Bullet")

    _add_heading(doc, "Top risks", 2)
    for rk in pkg["risks"][:8]:
        doc.add_paragraph(f"{rk['id']}: {rk['risk']} | Mitigation: {rk['mitigation']}", style="List Bullet")

    doc.save(out_path)
